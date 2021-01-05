from django.shortcuts import render
from django.core.mail import send_mail
from django.conf import settings
# Create your views here.
# Prerequisites

from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.template import loader
import requests
import random
# Models

import psycopg2
from django.contrib.auth.models import User, auth
from django.db.models import Q
from django.contrib import messages
from .models import Event
from .models import Messaging
from .models import Comment
from .models import Preevent
from .models import Precomment
from .models import Groupmessages
from .models import Groupmembers
from .models import Details
from .models import Addreport
from .models import Registration
from django.contrib.staticfiles import finders
from django.templatetags.static import static
from django.contrib.staticfiles.storage import staticfiles_storage
# Date related and others
import io
from django.http import FileResponse
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch, cm
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


from datetime import timedelta
from datetime import datetime,timedelta
from datetime import datetime as dt, timedelta
from json import dumps
from django.core.files import File

from twilio.rest import TwilioRestClient
from twilio.rest import Client

# Dataexports
import time
from pathlib import Path
from os import path
import pytz
import os
import fpdf
from fpdf import Template
from docx import Document
import xlwt
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from docx.shared import RGBColor

# ----------------------------------------------------------------------prerequisites-------------------------------------------------------------------------------
# ----------------------------------------------------------------------prerequisites-------------------------------------------------------------------------------
# ----------------------------------------------------------------------prerequisites-------------------------------------------------------------------------------
# ----------------------------------------------------------------------prerequisites-------------------------------------------------------------------------------
# ----------------------------------------------------------------------prerequisites-------------------------------------------------------------------------------
# ----------------------------------------------------------------------prerequisites-------------------------------------------------------------------------------
# ----------------------------------------------------------------------prerequisites-------------------------------------------------------------------------------

def index(request): 
    # create data dictionary 
    # dump data 
    return render(request, 'depactt/index.html')


def login(request): 
	# create data dictionary 
	# dump data 
	if request.method=='POST':
		uname=request.POST['usname']
		passwordrr1=request.POST['pssword']

		if (len(uname)==0) or (len(passwordrr1)==0):
			messages.warning(request,'Please Enter Password Correctly')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
		else:
			user = auth.authenticate(username=uname, password=passwordrr1)
			if user is not None:
				print('it worked')
				auth.login(request, user)
				# return user
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-home')
			else:
				messages.warning(request,'User doesnt exists or Password wrong')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
	else:
		return render(request, 'depactt/login.html')

def signup(request):
	if request.method=='POST':
		fname=request.POST['firstname']
		lname=request.POST['lastname']
		ename=request.POST['emailid']
		uname=request.POST['username']
		passwordrr1=request.POST['password1']
		passwordrr2=request.POST['password2']
		stat=request.POST['dd']
		if (len(fname)==0) or(len(lname)==0) or(len(ename)==0) or(len(uname)==0) or(len(passwordrr1)==0) or(len(passwordrr2)==0):
			messages.warning(request,'Please Enter All the Parameters Correctly')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
		else:
			if passwordrr1==passwordrr2:
				if User.objects.filter(username=uname).exists():
					messages.warning(request,'User already exists')
					return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
				else:
					if stat=='staff':
						passwordrr3=request.POST['password3']
						if passwordrr3=='kjsieitcompsteachers':
							user= User.objects.create_user(username=uname, first_name=fname, last_name=lname, email=ename, password=passwordrr1, is_staff=True)
							user.save()
							user2=Details(urref=User.objects.get(username=uname))
							user2.save()
							print('it worked');
							messages.success(request,'User Created Successfully, Please Login to continue')
						else:
							messages.warning(request,'Wrong Code')
					if stat=='student':
						user= User.objects.create_user(username=uname, first_name=fname, last_name=lname, email=ename, password=passwordrr1, is_staff=False)
						user.save()
						user2=Details(urref=User.objects.get(username=uname))
						user2.save()
						print('it worked');
						messages.success(request,'User Created Successfully, Please Login to continue')
					return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
			else: 
				messages.warning(request,'Passwords are not same')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')

	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')



# ------00000000000000000000000000000000000000----------------000000---------------0000000--0000000------------------0000000----------------------------
# ------00000000000000000000000000000000000000---------------00000000--------------0000000--000000000----------------0000000----------------------------
# ------00000000000000000000000000000000000000--------------0000000000-------------0000000--00000000000--------------0000000----------------------------
# ------0000000--------00000000--------0000000-------------000000000000------------0000000--0000000000000------------0000000----------------------------
# ------0000000--------00000000--------0000000------------000000-0000000-----------0000000--000000000000000----------0000000----------------------------
# ------0000000--------00000000--------0000000-----------0000000--0000000----------0000000--0000000--00000000--------0000000----------------------------
# ------0000000--------00000000--------0000000----------0000000----0000000---------0000000--0000000----00000000------0000000----------------------------
# ------0000000--------00000000--------0000000---------0000000------0000000--------0000000--0000000------00000000----0000000----------------------------
# ------0000000--------00000000--------0000000--------0000000--------0000000-------0000000--0000000--------00000000--0000000----------------------------
# ------0000000------------------------0000000-------0000000----------0000000------0000000--0000000----------000000000000000----------------------------
# ------0000000------------------------0000000------0000000------------0000000-----0000000--0000000------------0000000000000----------------------------
# ------0000000------------------------0000000-----0000000--------------0000000----0000000--0000000--------------00000000000----------------------------
# ------0000000------------------------0000000----0000000----------------0000000---0000000--0000000----------------000000000----------------------------
# ------0000000------------------------0000000---0000000------------------0000000--0000000--0000000------------------0000000----------------------------



def home(request):
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
	if not request.user.is_authenticated:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
	users=User.objects.all()
	dicevents = {
		"user_number": users
	}
	ist = pytz.timezone('Asia/Calcutta')
	c=datetime.now(ist)
	d=datetime.utcnow() + timedelta(days=-20)
	d2=datetime.utcnow() + timedelta(days=-3)
	f=datetime.utcnow() + timedelta(days=-3)
	f2=datetime.utcnow() + timedelta(days=-2)
	n=datetime.utcnow() + timedelta(days=3)
	gn=datetime.utcnow() + timedelta(days=-1)
	prof=Details.objects.all()
	events=Event.objects.filter(pubdate__gt=d)
	eventsup=Event.objects.filter(date__gte=c,date__lt=n)
	eventsfin=Event.objects.filter(date__gt=f,date__lt=c,tags=request.user)
	noofevents=Event.objects.filter(tags=request.user)
	curstat=User.objects.get(username=currentuser)
	comments=Comment.objects.filter(id=0)
	addreport=Addreport.objects.filter(id=0)
	payload = {"head": "Welcome!", "body": "You are now in the home page"}
	user=User.objects.get(id=currentuserid)
	for i in range(len(noofevents)):
		print(i)
		cur=noofevents[i].title
		comments=comments | Comment.objects.filter(ecref=Event.objects.get(title=cur),cdate__gt=f2)
		addreport=addreport | Addreport.objects.filter(erref=Event.objects.get(title=cur),rdate__gt=f)
	if request.user.is_superuser==True and request.user.is_staff==True:
		preevents=Preevent.objects.filter(pubdate__gt=d2)
		events2=Preevent.objects.filter(pubdate__gt=d)
		noofevents2=Preevent.objects.filter(tags=request.user)
		comments2=Precomment.objects.filter(id=0)
		for i in range(len(noofevents2)):
			print(i)
			cur=noofevents2[i].id
			comments2=comments2 | Precomment.objects.filter(ecref=Preevent.objects.get(id=cur),cdate__gt=f2)
		dicevents["pre_event"] = preevents
		dicevents["comment_num2"] = comments2
		dicevents["event_number2"] = events2
	if request.user.is_staff==False:
		certi=Event.objects.filter(certion=True)
		certieve2=Event.objects.filter(id=0)
		for i in range(len(certi)):
			cur=certi[i].id
			certieve=Event.objects.get(id=cur)
			reg=Registration.objects.filter(erref=certieve, rby=request.user)
			if len(reg)!=0:
				certieve2=certieve2 | Event.objects.filter(id=cur)
		dicevents["certificate"] = certieve2
	noofgroups=Groupmembers.objects.filter(members=request.user)
	curstat=User.objects.get(username=currentuser)
	groupm2=Groupmessages.objects.filter(id=0)
	print(groupm2)
	for i in range(len(noofgroups)):
		print(i)
		cur=noofgroups[i].name
		groupm2=groupm2 | Groupmessages.objects.filter(groupref=Groupmembers.objects.get(name=cur)).exclude(seen=request.user).exclude(gmessagesender=curstat)
		print(groupm2)
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
		messagesm=Messaging.objects.filter(receiver=request.user,messageseen=False)
		noofmessages=0
		inbox=[]
		inbox2=[]
		msender=[]
		for i in range(0,len(messagesm)):
			cur=messagesm[i]
			if cur.messageseen==False:
				if cur.messagesender in inbox:
					print('it exists')
				else:
					inbox.append(cur.messagesender)
		print(inbox)
		dicevents["nomessages"] = inbox
		dicevents["nomessages2"] = messagesm
		dicevents["event_number"] = events
		dicevents["profile_number"] = prof
		dicevents["event_upcoming"] = eventsup
		dicevents["event_finished"] = eventsfin
		dicevents["comment_num"] = comments
		dicevents["report_num"] = addreport
		dicevents["group_num"] = groupm2
	return render(request, 'depactt/home.html',dicevents)


def profile(request): 
	users=Details.objects.filter(urref=request.user)
	dicevents = {
		"user_number": users
	}

	return render(request, 'depactt/profile.html',dicevents)


def help(request): 

	return render(request, 'depactt/help.html')

def allevents(request): 
	events=Event.objects.all().order_by('-pubdate')
	dicevents = {
		"event_number": events
	}
	ist = pytz.timezone('Asia/Calcutta')
	c=datetime.now(ist)
	events2=Event.objects.filter(date__gte=c)
	events3=Event.objects.filter(date__lt=c)
	comments=Comment.objects.all().order_by('cdate')
	report=Addreport.objects.all()
	dicevents["event_number2"] = events2
	dicevents["event_number3"] = events3
	dicevents["comment_num"] = comments
	dicevents["report_num"] = report
	return render(request, 'depactt/allevents.html',dicevents)

def finishedevents(request): 
	ist = pytz.timezone('Asia/Calcutta')
	c=datetime.now(ist)
	events=Event.objects.filter(date__lt=c)
	dicevents = {
		"event_number": events
	}
	return render(request, 'depactt/finishedevents.html',dicevents)

def upcomingevents(request): 
	ist = pytz.timezone('Asia/Calcutta')
	c=datetime.now(ist)
	events=Event.objects.filter(date__gte=c)
	dicevents = {
		"event_number": events
	}
	return render(request, 'depactt/upcomingevents.html',dicevents)

def addevent(request):
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
	ist = pytz.timezone('Asia/Calcutta')
	c=datetime.now(ist)
	eventscal = Event.objects.filter(date__lt=c).order_by('date')
	eventscal2 = Event.objects.filter(date__gte=c).order_by('date')
	users=User.objects.order_by('username')
	dicevents = {
		"event_number": eventscal
	}
	dicevents["event_number2"] = eventscal2
	dicevents["user_num"] = users
	if request.method=='POST':
		eventpos=request.FILES['eventpos']
		eventname=request.POST['eventname']
		eventpresenter=request.POST['eventpresenter']
		eventdesignation=request.POST['eventdesignation']
		eventorgc=request.POST['eventcommittee']
		eventtype=request.POST['eventtype']
		eventda=request.POST['eventdate']
		eventfrotime=request.POST['eventtimefro']
		eventtotime=request.POST['eventtimeto']
		link=request.POST['regin']
		fromreg=request.POST['regfrom']
		toreg=request.POST['regto']
		whatsapp=request.POST['eventwhatsapp']
		eventcon1=request.POST['eventconvener1']
		eventcon2=request.POST['eventconvener2']
		eventcontact1=request.POST['contact1']
		eventcontact2=request.POST['contact2']
		eventorg1=request.POST['eventorganizer1']
		eventorg2=request.POST['eventorganizer2']
		eventorg3=request.POST['eventorganizer3']
		eventdep=request.POST['eventdepartment']
		eventdescription=request.POST['eventdesc']
		eventperks=request.POST['eventperk']
		gmem=request.POST['nooftag']
		eventfin=request.POST.get('eventcertificate',False)
		ecom=request.POST['com']
		publisher=request.user
		print(publisher.is_staff)
		if Event.objects.filter(title=eventname).exists():
			messages.warning(request,'Event already exists.')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-addevent')
		else:
			if (len(eventname)==0) or(len(eventda)==0) or(len(eventfrotime)==0) or(len(eventtotime)==0) or(len(fromreg)==0) or(len(toreg)==0):
				messages.warning(request,'Please Enter All the Parameters Correctly')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-addevent')
			uss=User.objects.get(username=currentuser)
			if uss.is_staff==True:
				eventt= Event(poster=eventpos,title=eventname,description=eventdescription,perks=eventperks,link=link,whatsapp=whatsapp, date=eventda,tfrom=eventfrotime,tto=eventtotime,regfrom=fromreg,regto=toreg,etype=eventtype,presenter=eventpresenter,presenterd=eventdesignation, organizer=eventorgc, teacher1=eventorg1, teacher2=eventorg2, teacher3=eventorg3, convener1=eventcon1, convener2=eventcon2, contact1=eventcontact1, contact2=eventcontact2, certi=eventfin, department=eventdep, by=publisher, reminders=ecom)
				eventt.save()
				gmem2=int(gmem)
				for gm in range(1,gmem2):
					gm2=str(gm)
					ele='tag'+gm2
					globals()['naming'] = request.POST[ele]
					print(naming)
					currentt=User.objects.get(username=naming)
					print(currentt)
					eventt.tags.add(currentt)
				messages.success(request,'Event Successfully added')
			else:
				messages.warning(request,'You are not permitted to add event')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-addevent')

	else:
		return render(request,'depactt/addevent.html',dicevents)


def preevent(request):
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
	ist = pytz.timezone('Asia/Calcutta')
	# c.datetime.datetime.strftime("%y/%m/%d")
	c=datetime.now(ist)
	eventscal = Event.objects.filter(date__lt=c).order_by('date')
	eventscal2 = Event.objects.filter(date__gte=c).order_by('date')
	eventscal3 = Preevent.objects.all().order_by('date')
	groupss=Groupmembers.objects.filter(members=request.user)
	comments=Precomment.objects.all().order_by('cdate')
	users=User.objects.order_by('username')
	dicevents = {
		"event_number": eventscal
	}
	dicevents["event_number2"] = eventscal2
	dicevents["event_number3"] = eventscal3
	dicevents["group_number"] = groupss
	dicevents["user_num"] = users
	dicevents["comment_num"] = comments
	if request.method=='POST':
		eventpresenter=request.POST['eventpresenter']
		eventdesignation=request.POST['eventdesignation']
		eventorgc=request.POST['eventcommittee']
		eventda=request.POST['eventdate']		
		eventfrotime=request.POST['eventtimefro']
		eventtotime=request.POST['eventtimeto']
		fromreg=request.POST['regfrom']
		toreg=request.POST['regto']
		eventdep=request.POST['eventdepartment']
		eventtar=request.POST['eventtarget']
		eventbud=request.POST['eventbudget']
		eventdescription=request.POST['eventdesc']
		gmem=request.POST['nooftag']
		publisher=request.user
		if (len(eventdescription)==0) or(len(eventda)==0) or(len(eventfrotime)==0) or(len(eventtotime)==0) or(len(fromreg)==0) or(len(toreg)==0):
			messages.warning(request,'Please Enter All the Parameters Correctly')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-addevent')
		uss=User.objects.get(username=currentuser)
		if uss.is_staff==True:
			eventt= Preevent(description=eventdescription, date=eventda,tfrom=eventfrotime,tto=eventtotime,regfrom=fromreg,regto=toreg,presenter=eventpresenter,presenterd=eventdesignation, organizer=eventorgc, department=eventdep, by=publisher,budget=eventbud,target=eventtar)
			eventt.save()
			gmem2=int(gmem)
			for gm in range(1,gmem2):
				gm2=str(gm)
				ele='tag'+gm2
				globals()['naming'] = request.POST[ele]
				print(naming)
				currentt=User.objects.get(username=naming)
				print(currentt)
				eventt.tags.add(currentt)
			messages.success(request,'Pre-Event Successfully added')
		else:
			messages.warning(request,'You are not permitted to add a pre-event')
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')

	else:
		return render(request,'depactt/preevent.html',dicevents)



def messaging(request): 
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
	users=User.objects.filter(is_staff=True).exclude(id=currentuserid)
	users2=User.objects.filter(is_staff=False).exclude(id=currentuserid)
	dicuser = {
		"user_number": users
	}
	dicuser["user_number2"]=users2
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
		messagesm=Messaging.objects.filter(receiver=request.user)
		noofmessages=0
		inbox=[]
		print(len(messagesm))
		for i in range(0,len(messagesm)):
			cur=messagesm[i]
			if cur.messageseen==False:
				if cur.messagesender in inbox:
					print('it exists')
				else:
					inbox.append(cur.messagesender)
		print(inbox)
		dicuser["nomessages"]=inbox
	return render(request, 'depactt/messages.html',dicuser)


def groups(request): 
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
	users2=User.objects.exclude(id=currentuserid).order_by('username')
	users=Groupmembers.objects.filter(members=request.user)
	noofgroups=Groupmembers.objects.filter(members=request.user)
	curstat=User.objects.get(username=currentuser)
	groupm2=Groupmessages.objects.filter(id=0)
	print(groupm2)
	for i in range(len(noofgroups)):
		print(i)
		cur=noofgroups[i].name
		groupm2=groupm2 | Groupmessages.objects.filter(groupref=Groupmembers.objects.get(name=cur)).exclude(seen=request.user).exclude(gmessagesender=curstat)
		print(groupm2)
	dicuser = {
		"user_number": users
	}
	dicuser["user_number2"]=users2
	dicuser["group_num"]=groupm2
	return render(request, 'depactt/groups.html',dicuser)




# ----------------------000000---------------00000000000000000000000-------00000000000000000000000------------------------------------
# ---------------------00000000--------------0000000000000000000000000-----0000000000000000000000000----------------------------------
# --------------------0000000000-------------00000000000000000000000000----00000000000000000000000000---------------------------------
# -------------------000000000000------------0000000-------------0000000---0000000-------------0000000--------------------------------
# ------------------000000-0000000-----------0000000-------------0000000---0000000-------------0000000--------------------------------
# -----------------0000000--0000000----------0000000-------------0000000---0000000-------------0000000--------------------------------
# ----------------0000000----0000000---------0000000-------------0000000---0000000-------------0000000--------------------------------
# ---------------0000000------0000000--------0000000-------------0000000---0000000-------------0000000--------------------------------
# --------------0000000--------0000000-------0000000-------------0000000---0000000-------------0000000--------------------------------
# -------------0000000----------0000000------0000000-------------0000000---0000000-------------0000000--------------------------------
# ------------0000000------------0000000-----0000000-------------0000000---0000000-------------0000000--------------------------------
# -----------0000000--------------0000000----00000000000000000000000000----00000000000000000000000000---------------------------------
# ----------0000000----------------0000000---0000000000000000000000000-----0000000000000000000000000----------------------------------
# ---------0000000------------------0000000--00000000000000000000000-------00000000000000000000000------------------------------------


def addgroup(request):
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
	if request.method=='POST':
		gname=request.POST['name']
		gmem=request.POST['members']
		parti2=request.POST['part1']
		parti3=request.POST['part2']
		parti4=request.POST['part3']
		parti5=request.POST['part4']
		parti6=request.POST['part5']
		parti7=request.POST['part6']
		parti8=request.POST['part7']
		parti9=request.POST['part8']
		parti10=request.POST['part9']
		parti11=request.POST['part10']
		if (len(parti2)==0) or(len(parti3)==0) or(len(gname)==0):
			messages.warning(request,'Please Enter All the Parameters Correctly')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-groups')
		else:
			if Groupmembers.objects.filter(name=gname).exists():
				messages.warning(request,'Group already exists')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-groups')
			else:
				if gmem:
					user= Groupmembers(name=gname, noofmembers=gmem, gby=request.user)
					user.save()
					user.members.add(request.user)
					gmem2=int(gmem)
					for gm in range(1,gmem2):
						gm2=str(gm)
						ele='part'+gm2
						globals()['naming'] = request.POST[ele]
						print(naming)
						currentt=User.objects.get(username=naming)
						print(currentt)
						user.members.add(currentt)
					messages.success(request,'Group Created Successfully, Be the first to message')
					return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-groups')
				else:
					messages.warning(request,'Group members should be from 3 to 10')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-groups')

	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-groups')


def comment(request):
	if request.method=='POST':
		eventname=request.POST['eventname']
		comment=request.POST['comment']
		publisher=request.user
		print(publisher)
		if len(comment)==0:
			messages.warning(request,'Please, type something first')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
		else:
			if not Event.objects.filter(title=eventname).exists():
				print('event exists')
				messages.warning(request,'Event doesnt exists.')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
			else:
				com= Comment(ecref=Event.objects.get(title = eventname), body=comment, cby=publisher)
				com.save()
				print('event doesnt exists or password wrong');
				messages.success(request,'Commented Successfully, Check the Post')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')

	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')



def precomment(request):
	if request.method=='POST':
		eventname=request.POST['eventname']
		comment=request.POST['comment']
		publisher=request.user
		print(publisher)
		if len(comment)==0:
			messages.warning(request,'Please, type something first')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')
		else:
			if not Preevent.objects.filter(id=eventname).exists():
				print('event exists')
				messages.warning(request,'Event doesnt exists.')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')
			else:
				com= Precomment(ecref=Preevent.objects.get(id = eventname), body=comment, cby=publisher)
				com.save()
				print('event doesnt exists or password wrong');
				messages.success(request,'Commented Successfully, Check the Post')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')

	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')


def addreport(request):
	if request.method=='POST':
		eventname=request.POST['eventname']
		reptitle=request.POST['reporttitle']
		repobj=request.POST['reportobj']
		repkep=request.POST['reportkep']
		repout=request.POST['reportout']
		repfed=request.POST['reportfed']
		certif=request.POST['reportcert']
		publisher=request.user
		print(publisher)
		if not Event.objects.filter(title=eventname).exists():
			print('event exists')
			messages.warning(request,'Event doesnt exists.')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
		else:
			rep= Addreport(erref=Event.objects.get(title = eventname),rtitle=reptitle, objectives=repobj, keypoints=repkep, outcomes=repout, feedback=repfed, rby=publisher)
			rep.save()
			if certif==True:
				cur=Event.objects.get(title = eventname)
				cur.certion=True
				cur.save()
				recipient_list = []
				users=Registration.objects.filter(erref=Event.objects.get(title=eventname))
				for i in range(len(users)):
					current=users[i].rby.username
					currentuser=User.objects.get(username=current)
					recipient_list.append(currentuser.email)
				subject = 'Certificate'
				st="\r\t"
				message = 'Certificate distribution for '+eventname+' is open,'+st+'Go and get your certificates on the website on event details on the Events panel,'+st+'Thank you.'+st+st+st+st+st+st+'If it was not you then Please contact the admin'
				email_from = settings.EMAIL_HOST_USER
				send_mail( subject, message, email_from, recipient_list )
			print('event doesnt exists or password wrong');
			messages.success(request,'Report Added Successfully')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')



def register(request):
	if request.method=='POST':
		eventname=request.POST['eventname']
		reptitle=request.POST['reporttitle']
		repbody=request.POST['reportbody']
		publisher=request.user
		print(publisher)
		if not Event.objects.filter(title=eventname).exists():
			print('event exists')
			messages.warning(request,'Event doesnt exists.')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
		else:
			rep= Registration(erref=Event.objects.get(title = eventname),name=reptitle, reason=repbody, rby=publisher)
			rep.save()
			ee= Event.objects.get(title = eventname)
			print('event doesnt exists or password wrong')
			cvv=str(ee.id)
			subject = 'Registration Successful ('+cvv+')'
			st="\r\t"
			message = publisher.first_name+','+st+'Your registration for '+eventname+' was successful,'+st+'Please be updated on event details on the Events panel,'+st+'Thank you.'+st+st+st+st+st+st+'If it was not you then Please contact the admin'
			email_from = settings.EMAIL_HOST_USER
			recipient_list = [publisher.email,]
			send_mail( subject, message, email_from, recipient_list )
			messages.success(request,'Registered Successfully')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')

def groupmessages(request): 
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
	users=Groupmembers.objects.filter(members=request.user)
	det=Details.objects.all()
	print(currentuserid)
	print(currentuser)
	print(users)
	dicuser = {
		"user_number": users
	}
	print(dicuser)
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("send", "")
		message=request.POST.get("message", "")
		sevent=request.POST.get("event", "")
		if len(message)!=0:
			if len(sevent)==0:
				mmsg=Groupmembers.objects.get(name=evid)
				mm=Groupmessages(gmessage=message,gmessagesender=request.user,groupref=mmsg)
				mm.save()
				messagesm=Groupmessages.objects.filter(groupref=mmsg).order_by('gmessagedate','gtime')
				dicuser["nomessages"]=messagesm
				current=Groupmembers.objects.filter(name=evid)
				dicuser["current_num"]=current
			else:
				mmsg=Groupmembers.objects.get(name=evid)
				mm=Groupmessages(gmessage=message,gmessagesender=request.user,preevent=Preevent.objects.get(id=sevent),groupref=mmsg,gmessageseen=True)
				mm.save()
				messagesm=Groupmessages.objects.filter(groupref=mmsg).order_by('gmessagedate','gtime')
				dicuser["nomessages"]=messagesm
				current=Groupmembers.objects.filter(name=evid)
				dicuser["current_num"]=current
		else:
			mmsg=Groupmembers.objects.get(name=evid)
			messagesm=Groupmessages.objects.filter(groupref=mmsg).order_by('gmessagedate','gtime')
			if Groupmessages.objects.filter(groupref=mmsg).exists():
				last=Groupmessages.objects.filter(groupref=mmsg).latest('gmessagedate','gtime')
				curstat=User.objects.get(username=currentuser)
				last2=Groupmessages.objects.filter(groupref=mmsg).exclude(gmessagesender=curstat)
				for i in range(len(last2)):
					last2[i].seen.add(curstat)

			noofgroups=Groupmembers.objects.filter(members=request.user)
			curstat=User.objects.get(username=currentuser)
			groupm2=Groupmessages.objects.filter(id=0)
			print(groupm2)
			for i in range(len(noofgroups)):
				print(i)
				cur=noofgroups[i].name
				groupm2=groupm2 | Groupmessages.objects.filter(groupref=Groupmembers.objects.get(name=cur)).exclude(seen=request.user).exclude(gmessagesender=curstat)
				print(groupm2)
			dicuser["nomessages"]=messagesm
			current=Groupmembers.objects.filter(name=evid)
			dicuser["current_num"]=current
			dicuser["group_num"]=groupm2
			dicuser["details"]=det
			if Groupmessages.objects.filter(groupref=mmsg, gmessagesender=request.user).exists():
				last2=Groupmessages.objects.filter(groupref=mmsg, gmessagesender=request.user).latest('gmessagedate','gtime')
				dicuser["latest"]=last2
		return render(request, 'depactt/groupmessages.html',dicuser)
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-groups')


def fullmessages(request): 
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
	users=User.objects.all()
	userlist = list(users)
	print(userlist)
	listnum=len(userlist)
	c=int(listnum)
	dicmessages = {
		"user_number": users
	}
	messagesbar= {

	}
	if request.method == "POST":
		print('entered post')
		sender=request.POST.get("send", "")
		evid=request.POST.get("usid", "")
		msgsender=request.POST.get("sender", "")
		receiver=request.POST.get("rece", "")
		message=request.POST.get("message", "")
		senderno=request.POST.get("sendno", "")
		print(evid)
		print(sender)
		print(currentuser)
		print(receiver)
		print(message)
		print(senderno)
		print(msgsender)
		if (len(receiver)==0) or(len(message)==0) or(len(senderno)==0) or(len(msgsender)==0):
			dicmessages["user_num"]=evid
			dicmessages["user_num2"]=sender
			messagesm=Messaging.objects.filter(Q(receiver=request.user,messagesenderid=User.objects.get(id=sender),messageseen=True) | Q(receiver=User.objects.get(username=evid),messagesenderid=User.objects.get(id=currentuserid))).order_by('messagedate','timee')
			dicmessages["msg"]=messagesm
			messagesms=Messaging.objects.filter(receiver=request.user,messagesenderid=User.objects.get(id=sender),messageseen=False)
			dicmessages["msg2"]=messagesms
			messagesm2=Messaging.objects.filter(receiver=request.user,messagesenderid=User.objects.get(id=sender))
			for i in range(0,len(messagesm2)):
				cur=messagesm2[i]
				print(messagesm2[i].messageseen)
				if cur.receiver==User.objects.get(username=currentuser):
					cur.messageseen=True
					cur.save()
		else :
			if User.objects.filter(username=receiver).exists():
				messagedone= Messaging(messagesender=User.objects.get(username=currentuser), receiver=User.objects.get(username=receiver), message=message, messagesenderid=User.objects.get(id=senderno))
				messagedone.save()
			dicmessages["user_num"]=evid
			dicmessages["user_num2"]=sender
			messagesm=Messaging.objects.filter(Q(receiver=User.objects.get(username=currentuser),messagesenderid=User.objects.get(id=sender)) | Q(receiver=User.objects.get(username=evid),messagesenderid=User.objects.get(id=currentuserid))).order_by('messagedate','timee')
			dicmessages["msg"]=messagesm
		return render(request, 'depactt/fullmessages.html',dicmessages)
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-messaging')

def addposter(request):
	if request.method=='POST':
		evid=request.POST.get("send","")
		posterurl=request.POST.get("poster","")
		event=Event.objects.get(id=evid)
		print(posterurl)
		path='posters/'+posterurl
		event.poster=posterurl
		event.save()
		print(event.title)
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')

# ----------00000000000000000000000--------00000000000000000000000---------0000000--00000000000000000000000-------------------------------
# ----------00000000000000000000000--------0000000000000000000000000-------0000000--00000000000000000000000-------------------------------
# ----------00000000000000000000000--------00000000000000000000000000------0000000--00000000000000000000000-------------------------------
# ----------0000000------------------------0000000-------------0000000-----0000000----------0000000---------------------------------------
# ----------0000000------------------------0000000-------------0000000-----0000000----------0000000---------------------------------------
# ----------0000000------------------------0000000-------------0000000-----0000000----------0000000---------------------------------------
# ----------00000000000000000000000--------0000000-------------0000000-----0000000----------0000000---------------------------------------
# ----------00000000000000000000000--------0000000-------------0000000-----0000000----------0000000---------------------------------------
# ----------0000000------------------------0000000-------------0000000-----0000000----------0000000---------------------------------------
# ----------0000000------------------------0000000-------------0000000-----0000000----------0000000---------------------------------------
# ----------0000000------------------------0000000-------------0000000-----0000000----------0000000---------------------------------------
# ----------00000000000000000000000--------00000000000000000000000000------0000000----------0000000---------------------------------------
# ----------00000000000000000000000--------0000000000000000000000000-------0000000----------0000000---------------------------------------
# ----------00000000000000000000000--------00000000000000000000000---------0000000----------0000000---------------------------------------


def grdel(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		print(evid)
		p=request.user.is_superuser
		print(p)
		if p==True:
			Groupmembers.objects.filter(id=evid).delete()
			messages.success(request,'Group Successfully deleted')
		else:
			messages.warning(request,'You cannot delete a group since you are not an admin')
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-groups')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-groups')


def alldel(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		print(evid)
		p=request.user.is_superuser
		print(p)
		if p==True:
			Event.objects.filter(id=evid).delete()
			messages.success(request,'Event Successfully deleted')
		else:
			messages.warning(request,'You cannot delete an event since you are not an admin')
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')



def predel(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		print(evid)
		p=request.user.is_superuser
		print(p)
		if p==True:
			Preevent.objects.filter(id=evid).delete()
			messages.success(request,'Event Successfully deleted')
		else:
			messages.warning(request,'You cannot delete an event since you are not an admin')
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')



def toadmin(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		evid2=request.POST.get("dell2", "")
		print(evid2)
		if evid2=='kjsomaiyaadminteachers':
			uss = User.objects.get(username = evid)
			uss.is_superuser=True
			uss.save()
			if request.user.is_superuser==True:
				messages.success(request,'You are now an admin')
			else:
				messages.warning(request,'Couldnt turn into admin')
		else:
			messages.warning(request,'Wrong Code')
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-profile')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-profile')

def tonormal(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		print(evid)
		uss = User.objects.get(username = evid)
		uss.is_superuser=False
		uss.save()
		if request.user.is_superuser==False:
			messages.warning(request,'You are not an admin now')
		else:
			messages.warning(request,'Couldnt turn into normal')
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-profile')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-profile')





def ongoing(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		event=Event.objects.get(id=evid)
		event.status='Ongoing'
		event.save()
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')

def confirmed(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		event=Event.objects.get(id=evid)
		event.status='Confirmed'
		event.save()
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')

def pending(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		event=Event.objects.get(id=evid)
		event.status='Pending'
		event.save()
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')


def toevent(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		evid2=request.POST.get("dell2", "")
		event=Preevent.objects.get(id=evid)
		mm=Event(title=evid2,description=event.description,target=event.target,date=event.date,tfrom=event.tfrom,tto=event.tto,regfrom=event.regfrom,regto=event.regto,presenter=event.presenter,presenterd=event.presenterd,organizer=event.organizer,department=event.department,by=request.user)
		mm.save()
		messages.success(request,'Converted to Event Successfully')
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')

def pongoing(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		event=Preevent.objects.get(id=evid)
		print(event.status)
		event.status='Ongoing'
		event.save()
		print(event.status)
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')

def pconfirmed(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		event=Preevent.objects.get(id=evid)
		print(event.status)
		event.status='Confirmed'
		event.save()
		print(event.status)
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')

def ppending(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		event=Preevent.objects.get(id=evid)
		print(event.status)
		event.status='Pending'
		event.save()
		print(event.status)
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-preevent')


def adddetails(request):
	if request.method == "POST":
		print('entered post')
		evid=request.POST.get("dell", "")
		sal=request.POST.get("usersal", "")
		desi=request.POST.get("userdes", "")
		cont1=request.POST.get("con1", "")
		cont2=request.POST.get("con2", "")
		colcod=request.POST.get("colcode", "")
		depa=request.POST.get("dep", "")
		Details.objects.filter(urref=request.user).update(salutation=sal, designation=desi,department=depa, contactinfo1=cont1, contactinfo2=cont2,color=colcod)
		print('done')
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-profile')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-profile')

# -----------000000000000000000000000000--0000000000000000000000000000--0000000----------0000000--00000000000000000000--00000000000000000000000000----------------------------------------
# -----------000000000000000000000000000--0000000000000000000000000000--0000000----------0000000--00000000000000000000--00000000000000000000000000----------------------------------------
# -----------000000000000000000000000000--0000000000000000000000000000--0000000----------0000000--0000000---------------00000000000000000000000000----------------------------------------
# -----------0000000-------------0000000------------00000000------------0000000----------0000000--0000000---------------0000000------------0000000----------------------------------------
# -----------0000000-------------0000000------------00000000------------0000000----------0000000--0000000---------------0000000------------0000000----------------------------------------
# -----------0000000-------------0000000------------00000000------------000000000000000000000000--00000000000000000000--0000000------------0000000----------------------------------------
# -----------0000000-------------0000000------------00000000------------000000000000000000000000--00000000000000000000--00000000000000000000000000----------------------------------------
# -----------0000000-------------0000000------------00000000------------000000000000000000000000--00000000000000000000--00000000000000000000000000----------------------------------------
# -----------0000000-------------0000000------------00000000------------000000000000000000000000--00000000000000000000--0000000-00000000---------------------------------------------------
# -----------0000000-------------0000000------------00000000------------0000000----------0000000--0000000---------------0000000---00000000----------------------------------------
# -----------0000000-------------0000000------------00000000------------0000000----------0000000--0000000---------------0000000-----00000000----------------------------------------
# -----------000000000000000000000000000------------00000000------------0000000----------0000000--0000000---------------0000000-------00000000----------------------------------------
# -----------000000000000000000000000000------------00000000------------0000000----------0000000--00000000000000000000--0000000---------00000000----------------------------------------
# -----------000000000000000000000000000------------00000000------------0000000----------0000000--00000000000000000000--0000000-----------00000000----------------------------------------



def userdetails(request): 
	if request.method == "POST":
		print('entered post')
		sender=request.POST.get("send", "")
		evid=request.POST.get("evid", "")
		print(sender)
		if User.objects.filter(username=sender).exists():
			user=User.objects.filter(username=sender)
			current=User.objects.get(username=sender)
			userdetails=Details.objects.filter(urref=current)
			dicuser = {
				"user_number2": user
			}
			dicuser["page"]=evid
			dicuser["details"]=userdetails
		else:
			print('not found')
		return render(request, 'depactt/userdetails.html',dicuser)
	else:
		if request.user.is_authenticated:
			currentuser = request.user.username
			currentuserid = request.user.id
		users=User.objects.exclude(id=currentuserid)
		print(currentuserid)
		print(currentuser)
		print(users)
		dicuser = {
			"user_number": users
		}
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-messaging')


def eventdetails(request): 
	if request.method == "POST":
		print('entered post')
		sender=request.POST.get("send", "")
		evid=request.POST.get("evid", "")
		combod=request.POST.get("comment", "")
		evna=request.POST.get("eventname", "")
		print(sender)
		if Event.objects.filter(id=sender).exists():
			if (len(combod)==0):
				user=Event.objects.filter(id=sender)
				userob=Event.objects.get(id=sender)
				comments=Comment.objects.filter(ecref=userob).order_by('cdate')
				addreport=Addreport.objects.filter(erref=Event.objects.get(id=sender))
				addreport2=Registration.objects.filter(erref=Event.objects.get(id=sender),rby=request.user)
				addreport3=Registration.objects.filter(erref=Event.objects.get(id=sender))
				print(addreport2)
				regiss=len(addreport3)
				dicuser = {
					"user_number2": user
				}
				dicuser["page"]=evid
				dicuser["comment_num"] = comments
				dicuser["report_num"] = addreport
				dicuser["regis_num"] = addreport2
				dicuser["regs"] = regiss
			else:
				user=Event.objects.filter(id=sender)
				userob=Event.objects.get(id=sender)
				com= Comment(ecref=userob, body=combod, cby=request.user)
				com.save()
				comments=Comment.objects.filter(ecref=userob).order_by('cdate')
				addreport=Addreport.objects.filter(erref=Event.objects.get(id=sender))
				dicuser = {
					"user_number2": user
				}
				dicuser["page"]=evid
				dicuser["comment_num"] = comments
				dicuser["report_num"] = addreport
			return render(request, 'depactt/eventdetails.html',dicuser)
		else:
			print('not found')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		if request.user.is_authenticated:
			currentuser = request.user.username
			currentuserid = request.user.id
		users=User.objects.exclude(id=currentuserid)
		print(currentuserid)
		print(currentuser)
		print(users)
		dicuser = {
			"user_number": users
		}
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')



def forgotpassword(request): 
	if request.method == "POST":
		print('entered post')
		user=request.POST.get("usname", "")
		print(user)
		if User.objects.filter(username=user).exists():
			if (len(user)!=0):
				users=User.objects.filter(username=user)
				dicuser = {
					"user_number": users
				}
				randnum=random.randint(000000,999999)
				curuser=User.objects.get(username=user)
				subject = 'Forgot Password'
				randnum2=str(randnum)
				message = ' Your Verification Code is : '+randnum2+', Please Do not share with anyone'
				email_from = settings.EMAIL_HOST_USER
				recipient_list = ['projectdepactmanagement@gmail.com',curuser.email,]
				send_mail( subject, message, email_from, recipient_list )
				print(curuser.username)
				print(curuser.password)
				print(curuser.email)
				print(curuser.first_name)
				dicuser["ver_code"] = randnum2
				messages.success(request,'A Mail Sent to your registered email')
				return render(request, 'depactt/forgotpassword1.html',dicuser)
			else:
				messages.warning(request,'Please Enter a correct username')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
		else:
			print('not found')
			messages.warning(request,'Please Enter a correct username')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
	else:
		if request.user.is_authenticated:
			currentuser = request.user.username
			currentuserid = request.user.id
		users=User.objects.exclude(id=currentuserid)
		print(currentuserid)
		print(currentuser)
		print(users)
		dicuser = {
			"user_number": users
		}
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')


def forgotpassword2(request): 
	if request.method == "POST":
		print('entered post')
		user=request.POST.get("uname", "")
		vcode1=request.POST.get("vcode1", "")
		vcode2=request.POST.get("vcode2", "")
		pass1=request.POST.get("pass1", "")
		pass2=request.POST.get("pass2", "")
		print(user)
		if User.objects.filter(username=user).exists():
			if (len(pass1)!=0):
				users=User.objects.filter(username=user)
				dicuser = {
					"user_number": users
				}
				if pass1==pass2:
					if vcode1==vcode2:
						curuser=User.objects.get(username=user)
						passmain=str(pass1)
						curuser.set_password(passmain)
						curuser.save()
						print(curuser.username)
						print(curuser.password)
						print(curuser.email)
						print(curuser.first_name)
						print('Sent vcode : '+vcode1)
						print('Entered vcode : '+vcode2)
						print(pass1)
						print(pass2)
						subject = 'SUCCESS'
						message = 'Password has been reset successfully, '+curuser.first_name+'. If it was not you please contact the admin'
						email_from = settings.EMAIL_HOST_USER
						recipient_list = ['projectdepactmanagement@gmail.com',curuser.email,]
						send_mail( subject, message, email_from, recipient_list )
						messages.success(request,'Password successfully reset')
					else:
						messages.success(request,'Wrong Verification Code')
				else:
					messages.success(request,'Passwords are not the same')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
			else:
				messages.warning(request,'Please Enter a password first')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
		else:
			print('not found')
			messages.warning(request,'User not found')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')
	else:
		if request.user.is_authenticated:
			currentuser = request.user.username
			currentuserid = request.user.id
		users=User.objects.exclude(id=currentuserid)
		print(currentuserid)
		print(currentuser)
		print(users)
		dicuser = {
			"user_number": users
		}
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')


def eventdetails2(request): 
	if request.method == "POST":
		print('entered post')
		sender=request.POST.get("dell", "")
		evid=request.POST.get("evid", "")
		print(sender)
		if Event.objects.filter(id=sender).exists():
			user=Event.objects.filter(id=sender)
			userob=Event.objects.get(id=sender)
			dicuser = {
				"user_number2": user
			}
			dicuser["page"]=evid
			return render(request, 'depactt/eventdetails2.html',dicuser)
		else:
			print('not found')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		if request.user.is_authenticated:
			currentuser = request.user.username
			currentuserid = request.user.id
		users=User.objects.exclude(id=currentuserid)
		print(currentuserid)
		print(currentuser)
		print(users)
		dicuser = {
			"user_number": users
		}
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')


def editevent(request):
	if request.user.is_authenticated:
		currentuser = request.user.username
		currentuserid = request.user.id
	ist = pytz.timezone('Asia/Calcutta')
	c=datetime.now(ist)
	eventscal = Event.objects.filter(date__lt=c).order_by('date')
	eventscal2 = Event.objects.filter(date__gte=c).order_by('date')
	users=User.objects.order_by('username')
	dicevents = {
		"event_number": eventscal
	}
	dicevents["event_number2"] = eventscal2
	dicevents["user_num"] = users
	if request.method=='POST':
		eventname=request.POST['eventname']
		eventpresenter=request.POST['eventpresenter']
		eventdesignation=request.POST['eventdesignation']
		eventorgc=request.POST['eventcommittee']
		eventtype=request.POST['eventtype']
		eventda=request.POST['eventdate']
		eventfrotime=request.POST['eventtimefro']
		eventtotime=request.POST['eventtimeto']
		link=request.POST['regin']
		fromreg=request.POST['regfrom']
		toreg=request.POST['regto']
		whatsapp=request.POST['eventwhatsapp']
		eventcon1=request.POST['eventconvener1']
		eventcon2=request.POST['eventconvener2']
		eventcontact1=request.POST['contact1']
		eventcontact2=request.POST['contact2']
		eventorg1=request.POST['eventorganizer1']
		eventorg2=request.POST['eventorganizer2']
		eventorg3=request.POST['eventorganizer3']
		eventdep=request.POST['eventdepartment']
		eventdescription=request.POST['eventdesc']
		eventperks=request.POST['eventperk']
		eventfin=request.POST.get('eventcertificate',False)
		ecom=request.POST['com']
		publisher=request.user
		print(publisher.is_staff)
		if Event.objects.filter(title=eventname).exists():
			curevent=Event.objects.get(title=eventname)
			curevent.title=eventname
			curevent.presenter=eventpresenter
			curevent.presenterd=eventdesignation
			curevent.organizer=eventorgc
			curevent.etype=eventtype
			curevent.date=eventda
			curevent.tfrom=eventfrotime
			curevent.tto=eventtotime
			curevent.link=link
			curevent.regfrom=fromreg
			curevent.regto=toreg
			curevent.whatsapp=whatsapp
			curevent.convener1=eventcon1
			curevent.convener2=eventcon2
			curevent.contact1=eventcontact1
			curevent.contact2=eventcontact2
			curevent.teacher1=eventorg1
			curevent.teacher2=eventorg2
			curevent.teacher3=eventorg3
			curevent.department=eventdep=request.POST['eventdepartment']
			curevent.description=eventdescription
			curevent.perks=eventperks
			curevent.save()
			messages.success(request,'Event successfully edited')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
		else:
			if (len(eventname)==0) or(len(eventda)==0) or(len(eventfrotime)==0) or(len(eventtotime)==0) or(len(fromreg)==0) or(len(toreg)==0):
				messages.warning(request,'Please Enter All the Parameters Correctly')
				return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
			uss=User.objects.get(username=currentuser)
			if uss.is_staff==True:
				eventt= Event(title=eventname,description=eventdescription,perks=eventperks,link=link,whatsapp=whatsapp, date=eventda,tfrom=eventfrotime,tto=eventtotime,regfrom=fromreg,regto=toreg,etype=eventtype,presenter=eventpresenter,presenterd=eventdesignation, organizer=eventorgc, teacher1=eventorg1, teacher2=eventorg2, teacher3=eventorg3, convener1=eventcon1, convener2=eventcon2, contact1=eventcontact1, contact2=eventcontact2, certi=eventfin, department=eventdep, by=publisher, reminders=ecom)
				eventt.save()
				messages.success(request,'Event Successfully added')
			else:
				messages.warning(request,'You are not permitted to add event')
			return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')

	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')

# ----------0000000------------------------000000000000000000000000000---------00000000000000000000000000----------------------------------------
# ----------0000000------------------------000000000000000000000000000---------00000000000000000000000000----------------------------------------
# ----------0000000------------------------000000000000000000000000000---------00000000000000000000000000----------------------------------------
# ----------0000000------------------------0000000-------------0000000---------0000000-----------------------------------------------------------
# ----------0000000------------------------0000000-------------0000000---------0000000-----------------------------------------------------------
# ----------0000000------------------------0000000-------------0000000---------0000000-----------------------------------------------------------
# ----------0000000------------------------0000000-------------0000000---------0000000-----00000000000000----------------------------------------
# ----------0000000------------------------0000000-------------0000000---------0000000-----00000000000000----------------------------------------
# ----------0000000------------------------0000000-------------0000000---------0000000------------0000000----------------------------------------
# ----------0000000------------------------0000000-------------0000000---------0000000------------0000000----------------------------------------
# ----------0000000------------------------0000000-------------0000000---------0000000------------0000000----------------------------------------
# ----------00000000000000000000000--------000000000000000000000000000---------00000000000000000000000000----------------------------------------
# ----------00000000000000000000000--------000000000000000000000000000---------00000000000000000000000000----------------------------------------
# ----------00000000000000000000000--------000000000000000000000000000---------00000000000000000000000000----------------------------------------

def homelogout(request):
	auth.logout(request)
	return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-login')

# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------
# ---------------------------this is where data transfer starts---------------------------------



def toxl(request):
	response = HttpResponse(
		content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
	)
	response['Content-Disposition'] = 'attachment; filename=Events.xlsx'.format(
		date=datetime.now().strftime('%Y-%m-%d'),
	)
	wb=Workbook()
	ws=wb.active
	ws.title='Events'
	row_num = 1
	columns=['Event Name','Event type','Date','Time From','Time To','Description','Organizer','Department','Presnter','Presenter Designation','Faculty Co-ordinator 1','Faculty Co-ordinator 2','Faculty Co-ordinator 3','Convener 1','Convener 2','Contact','Published date','Published by']
	print(len(columns))
	for col_num, column_title in enumerate(columns, 1):
		cell = ws.cell(row=row_num, column=col_num)
		cell.value = column_title
	width= len("Faculty Co-ordinator 1")
	column_widths = []
	rows = Event.objects.all()
	for ro in rows:
		row_num += 1
		row = [ro.title,ro.etype,ro.date,ro.tfrom,ro.tto,ro.description,ro.organizer,ro.department,ro.presenter,ro.presenterd,ro.teacher1,ro.teacher2,ro.teacher3,ro.convener1,ro.convener2,ro.contact1,ro.pubdate,ro.by.username]
		for col_num, cell_value in enumerate(row, 1):
			cell = ws.cell(row=row_num, column=col_num)
			cell.value = str(cell_value)
			print(type(cell_value))
			if len(column_widths) > col_num:
				if len(str(cell_value)) > column_widths[col_num]:
					column_widths[col_num] = len(str(cell_value))
			else:
				column_widths += [len(str(cell_value))]
	for i, column_width in enumerate(column_widths):
		ws.column_dimensions[get_column_letter(i+1)].width = column_width
	if path.isfile('Downloads\\Events.xlsx')==True:
		messages.warning(request,"Events.xlsx already exists in your desktop")
	else:
		wb.save(response)
		messages.success(request,"Downloaded data to Events.xlsx file in your desktop")
	return response



def toreg(request):
	if request.method == "POST":
		evid=request.POST.get("dell", "")
		evid2=str(evid)
		response = HttpResponse(
			content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
		)
		response['Content-Disposition'] = 'attachment; filename=Registrations-{time}.xlsx'.format(
			time = datetime.now().strftime("%H:%M:%S"),
		)
		wb=Workbook()
		ws=wb.active
		ws.title='Events'
		row_num = 1
		evid2=str(evid)
		columns=['Event Name','Event type','Date','Time From','Time To','Student Name','Reason','Student Username','Student Email']
		print(len(columns))
		for col_num, column_title in enumerate(columns, 1):
			cell = ws.cell(row=row_num, column=col_num)
			cell.value = column_title
		width= len("Faculty Co-ordinator 1")
		column_widths = []
		rows = Registration.objects.filter(erref=Event.objects.get(id=evid))
		for ro in rows:
			row_num += 1
			row = [ro.erref.title,ro.erref.etype,ro.erref.date,ro.erref.tfrom,ro.erref.tto,ro.name,ro.reason,ro.rby.username,ro.rby.email]
			for col_num, cell_value in enumerate(row, 1):
				cell = ws.cell(row=row_num, column=col_num)
				cell.value = str(cell_value)
				print(type(cell_value))
				if len(column_widths) > col_num:
					if len(str(cell_value)) > column_widths[col_num]:
						column_widths[col_num] = len(str(cell_value))
				else:
					column_widths += [len(str(cell_value))]
		for i, column_width in enumerate(column_widths):
			ws.column_dimensions[get_column_letter(i+1)].width = column_width
		if path.isfile('Registrations'+evid2+'.xlsx')==True:
			messages.warning(request,"Registrations"+evid2+".xlsx already exists in your desktop")
		else:
			wb.save(response)
			messages.success(request,"Downloaded data to Registrations"+evid2+".xlsx file in your desktop")
		return response
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')



def totxtall(request):
	if request.method == "POST":
		evid=request.POST.get("dell", "")
		title=request.POST.get("delltitle", "")
		date=request.POST.get("delldate", "")
		desc=request.POST.get("delldesc", "")
		etype=request.POST.get("delltype", "")
		timef=request.POST.get("dellfrom", "")
		timet=request.POST.get("dellto", "")
		regf=request.POST.get("dellregf", "")
		regt=request.POST.get("dellregt", "")
		org=request.POST.get("dellorg", "")
		pre=request.POST.get("dellpresenter", "")
		pred=request.POST.get("dellpresenterd", "")
		dep=request.POST.get("delldep", "")
		per=request.POST.get("dellperks", "")
		cer=request.POST.get("dellcerti", "")
		faccod1=request.POST.get("dellfaccod1", "")
		faccod2=request.POST.get("dellfaccod2", "")
		faccod3=request.POST.get("dellfaccod3", "")
		conv1=request.POST.get("dellconv1", "")
		conv2=request.POST.get("dellconv2", "")
		cont1=request.POST.get("dellcont1", "")
		cont2=request.POST.get("dellcont2", "")
		file='Event'+evid+'data.txt'
		response = HttpResponse(content_type='text/txt')
		content = "attachment; filename=%s " %(file)
		response['Content-Disposition'] = content
		response.write(""+title+"\t")
		response.write("Date : "+date+"\t")
		response.write("From "+timef+"\t")
		response.write("To "+timet+"\t")
		response.write(" ""\t")
		response.write("Description : "+desc+"\t")
		response.write("Presenter : "+pre+"\t")
		response.write(", "+pred+"\t")
		response.write(" ""\t")
		response.write("Organizer : "+org+"\t")
		response.write("Department : "+dep+"\t")
		response.write(" ""\t")
		response.write("Registrations from "+regf+" "+regt+"\t")
		response.write(" ""\t")
		response.write("Faculty Co-ordinators : "+faccod1+" ")
		response.write(" "+faccod2+" ")
		response.write(" "+faccod3+"\t")
		response.write("Conveners : "+conv1+" ")
		response.write(""+conv2+"\t")
		response.write("Contacts : "+cont1+" ")
		response.write(""+cont2+"\t")
		messages.success(request,"Downloaded data to Event"+evid+"data.txt file in your desktop")
		return response
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')



def topdfall(request):
	if request.method == "POST":
		evid=request.POST.get("dell", "")
		title=request.POST.get("delltitle", "")
		date=request.POST.get("delldate", "")
		desc=request.POST.get("delldesc", "")
		etype=request.POST.get("delltype", "")
		timef=request.POST.get("dellfrom", "")
		timet=request.POST.get("dellto", "")
		regf=request.POST.get("dellregf", "")
		regt=request.POST.get("dellregt", "")
		org=request.POST.get("dellorg", "")
		pre=request.POST.get("dellpresenter", "")
		pred=request.POST.get("dellpresenterd", "")
		dep=request.POST.get("delldep", "")
		per=request.POST.get("dellperks", "")
		cer=request.POST.get("dellcerti", "")
		faccod1=request.POST.get("dellfaccod1", "")
		faccod2=request.POST.get("dellfaccod2", "")
		faccod3=request.POST.get("dellfaccod3", "")
		conv1=request.POST.get("dellconv1", "")
		conv2=request.POST.get("dellconv2", "")
		cont1=request.POST.get("dellcont1", "")
		cont2=request.POST.get("dellcont2", "")
		pdf = fpdf.FPDF(format='letter')
		pdf.add_page()
		pdf.set_text_color(167, 36, 41)
		pdf.set_fill_color(0, 250, 250)
		pdf.set_font('Arial', '', 14)
		pdf.set_font('Arial', '', 13)
		name1='https://departmentalactivity.herokuapp.com/static/kj.png'
		name2='https://departmentalactivity.herokuapp.com/static/kj2.png'
		font1='https://departmentalactivity.herokuapp.com/static/algerianr.ttf'
		font2='//db.onlinewebfonts.com/c/527c5ab608cab860a6aae8ce02e14b0e?family=Calibri'
		pdf.cell(194,2,' ',0,1,'C')
		pdf.cell(194,5.5,'K.J. Somaiya Insititute of Engineering and',0,1,'C')
		pdf.image(name1, x = 20, y = 10, w = 24, h = 24, type = 'png', link = 'static')
		pdf.set_fill_color(0, 0, 250)
		pdf.cell(194,5.5,'Information Technology Sion, Mumbai',0,1,'C')
		pdf.cell(194,1,' ',0,1,'C')
		pdf.set_font('Times', 'B', 8)
		pdf.set_text_color(3, 90, 171)
		pdf.cell(194,3,'KJSIEIT IS ACCREDITED BY NAAC WITH "A" GRADE',0,1,'C')
		pdf.cell(194,3,'THREE PROGRAMS ACCREDITED BY NATIONAL BOARD OF ACCREDITATION',0,1,'C')
		pdf.cell(194,3,'BEST COLLEGE AWARDED BY UNIVERSITY OF MUMBAI, ISTE (MH) AND CSI (MUMBAI)',0,1,'C')
		pdf.image(name2, x = 175, y = 13, w = 24, h = 18, type = 'png', link = 'static')
		pdf.ln()
		pdf.line(10, 45, 205, 45)
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.set_font('Arial', '', 14)
		pdf.set_text_color(0,0,0)
		if org:
			pdf.cell(195,4,org,0,1,'C')
			pdf.ln()
			pdf.set_font('Arial', '', 11)
			pdf.set_text_color(100,100,100)
			pdf.cell(195,4,'is organizing a '+etype+' on topic',0,1,'C')
			pdf.ln()
		pdf.set_font("Arial", size=22)
		pdf.set_text_color(0,0,0)
		pdf.cell(195,11,title,0,1,'C')
		pdf.set_font('Arial', '', 11)
		pdf.set_text_color(100,100,100)
		if dep:
			pdf.cell(195,4,'in association with',0,1,'C')
			pdf.ln()
			pdf.set_font('Arial', '', 14)
			pdf.set_text_color(0,0,0)
			pdf.cell(195,4,dep,0,1,'C')
		pdf.set_font("Arial", size=12)
		pdf.set_text_color(30,30,30)
		pdf.cell(195,6,'Date : '+date+' ',0,1,'R')
		pdf.cell(195,6,'Time : '+timef+' - '+timet,0,1,'R')
		pdf.set_font('Arial', '', 11)
		pdf.set_text_color(100,100,100)
		if regf:
			pdf.cell(195,4,'Registrations from '+regf+' to '+regt,0,1,'L')
		pdf.ln()
		pdf.set_font('Arial', 'B', 11)
		pdf.set_text_color(0,0,0)
		if desc:
			pdf.cell(195,4,'A short Description : ',0,1,'L')
			pdf.ln()
			pdf.set_font('Arial', '', 11)
			pdf.cell(195,4,desc,0,1,'L')
		else:
			pdf.ln()
			pdf.ln()
			pdf.ln()
		pdf.ln()
		pdf.set_font('Arial', 'B', 11)
		pdf.set_text_color(0,0,0)
		if pre:
			pdf.cell(195,4,'Description about Presenter',0,1,'R')
			pdf.ln()
			pdf.set_font('Arial', '', 11)
			pdf.cell(195,4,pre,0,1,'R')
			pdf.ln()
			pdf.cell(195,4,pred,0,1,'R')
			pdf.ln()
		if per:
			pdf.set_font('Arial', 'B', 11)
			pdf.set_text_color(0,0,0)
			pdf.cell(195,4,'Perks/Benefits  : ',0,1,'L')
			pdf.ln()
			pdf.set_font('Arial', '', 11)
			pdf.cell(195,4,per,0,1,'L')
			pdf.ln()
		else:
			pdf.ln()
			pdf.ln()
			pdf.ln()
			pdf.ln()
		if cer=='True':
			pdf.cell(195,4,'Certificates will be given ',0,1,'L')
			pdf.ln()
		else:
			pdf.ln()
			pdf.ln()
		pdf.line(10, 230, 205, 230)
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.set_font('Arial', '', 11)
		pdf.set_text_color(0,0,0)
		pdf.cell(195,4,"Faculty Co-ordinators :  "+faccod1+",  "+faccod2+",  "+faccod3,0,1,'L')
		pdf.ln()
		pdf.cell(195,4,"Conveners :  "+conv1+",  "+conv2,0,1,'L')
		pdf.ln()
		pdf.cell(195,4,"Contact :  "+cont1+",  "+cont2,0,1,'L')
		file='Event'+evid+'_data.pdf'
		response = HttpResponse(pdf.output(name='Event'+evid+'_data.pdf',dest='S').encode('latin-1'), content_type='application/pdf')
		content = "attachment; filename=%s " %(file)
		response['Content-Disposition'] = content
		messages.success(request,"Downloaded Report to Event"+evid+"_data.pdf file in your desktop")
		return response



def tocerti(request):
	if request.method == "POST":
		evid=request.POST.get("dell", "")
		etype=request.POST.get("delltype", "")
		title=request.POST.get("delltitle", "")
		date=request.POST.get("delldate", "")
		org=request.POST.get("dellorg", "")
		dep=request.POST.get("delldep", "")
		con1=request.POST.get("dellconv1", "")
		con2=request.POST.get("dellconv2", "")
		pdf = fpdf.FPDF(format='letter')
		pdf = fpdf.FPDF('L', 'mm', (215.9, 279.4))
		pdf.add_page()
		pdf.set_font('Arial', '', 13)
		pdf.set_text_color(167, 36, 41)
		pdf.set_fill_color(0, 250, 250)
		name1='https://departmentalactivity.herokuapp.com/static/kj.png'
		name2='https://departmentalactivity.herokuapp.com/static/kj2.png'
		font1='https://departmentalactivity.herokuapp.com/static/algerianr.ttf'
		font2='https://departmentalactivity.herokuapp.com/static/calibb.ttf'
		pdf.cell(279.4,22,'',0,1,'C')
		pdf.cell(259.4,5.5,'K.J. Somaiya Insititute of Engineering and',0,1,'C')
		pdf.image(name1, x = 50, y = 30, w = 24, h = 24, type = 'png', link = 'static')
		pdf.set_fill_color(0, 0, 250)
		pdf.cell(259.4,5.5,'Information Technology Sion, Mumbai',0,1,'C')
		pdf.cell(259.4,1,'',0,1,'C')
		pdf.set_font('Times', 'B', 8)
		pdf.set_text_color(3, 90, 171)
		pdf.cell(259.4,3,'KJSIEIT IS ACCREDITED BY NAAC WITH "A" GRADE',0,1,'C')
		pdf.cell(259.4,3,'THREE PROGRAMS ACCREDITED BY NATIONAL BOARD OF ACCREDITATION',0,1,'C')
		pdf.cell(259.4,3,'BEST COLLEGE AWARDED BY UNIVERSITY OF MUMBAI, ISTE (MH) AND CSI (MUMBAI)',0,1,'C')
		pdf.image(name2, x = 210, y = 32, w = 24, h = 18, type = 'png', link = 'static')
		pdf.ln()
		# first
		pdf.line(25.4, 22, 254, 22)
		pdf.line(25.4, 193.9, 254, 193.9)
		pdf.line(254, 22, 254, 193.9)
		pdf.line(25.4, 22, 25.4, 193.9)
		# third
		pdf.line(33.4, 26, 246, 26) #top
		pdf.line(33.4, 189.9, 246, 189.9) #bottom
		pdf.line(250, 30, 250, 185.9) #right
		pdf.line(29.4, 30, 29.4, 185.9) #left
		# second
		pdf.line(23.4, 20, 23.4, 195.9)
		pdf.line(23.4, 20, 256, 20)
		pdf.line(23.4, 195.9, 256, 195.9)
		pdf.line(256, 20, 256, 195.9)
		# otherstopleftside
		pdf.line(33.4, 26, 33.4, 20) #goesup
		pdf.line(23.4, 30, 29.4, 30) #goesleft
		pdf.line(33.4, 20,29.4, 20) #leftfromtop
		pdf.line(23.4, 26, 23.4, 30) #topfromleft
		pdf.line(29.4, 26, 29.4, 20) #bottomfromleftfromtop
		pdf.line(29.4, 26, 23.4, 26) #rightfromtopfromleft
		# otherstoprightside
		pdf.line(246, 26, 246, 20) #goesup
		pdf.line(250, 30, 256, 30) #goesleft
		pdf.line(250, 20,246, 20) #leftfromtop
		pdf.line(256, 26, 256, 30) #topfromleft
		pdf.line(250, 26, 250, 20) #bottomfromleftfromtop
		pdf.line(250, 26, 256, 26) #rightfromtopfromleft

		# othersbottomleftside
		pdf.line(33.4, 189.9, 33.4, 195.9) #goesup
		pdf.line(23.4, 185.9, 29.4, 185.9) #goesleft
		pdf.line(33.4, 195.9,29.4, 195.9) #leftfromtop
		pdf.line(23.4, 189.9, 23.4, 185.9) #topfromleft
		pdf.line(29.4, 189.9, 29.4, 195.9) #bottomfromleftfromtop
		pdf.line(29.4, 189.9, 23.4, 189.9) #rightfromtopfromleft
		# othersbottomrightside
		pdf.line(246, 189.9, 246, 195.9) #goesup
		pdf.line(250, 185.9, 256, 185.9) #goesleft
		pdf.line(250, 195.9,246, 195.9) #leftfromtop
		pdf.line(256, 189.9, 256, 185.9) #topfromleft
		pdf.line(250, 189.9, 250, 195.9) #bottomfromleftfromtop
		pdf.line(250, 189.9, 256, 189.9) #rightfromtopfromleft
		pdf.ln()
		pdf.ln()
		pdf.set_font('Arial', '', 24)
		pdf.set_text_color(0,0,0)
		pdf.cell(259.4,10,' ',0,1,'C')
		pdf.cell(259.4,3,'CERTIFICATE OF PARTICIPATION',0,1,'C')
		currentfirst= request.user.first_name
		currentlast= request.user.last_name
		currentname=currentfirst+" "+currentlast
		pdf.set_font('Times', '', 12)
		pdf.set_text_color(50,50,50)
		pdf.cell(259.4,12,' ',0,1,'C')
		pdf.cell(259.4,3,'                                                                        This is to certify that        '+currentname,0,1,'L')
		pdf.line(130, 91, 190, 91)
		pdf.cell(259.4,1.5,' ',0,1,'C')
		pdf.cell(259.4,1.5,' ',0,1,'C')
		pdf.cell(259.4,1.5,' ',0,1,'C')
		pdf.cell(259.4,3,'has successfully completed '+etype+' on topic',0,1,'C')
		pdf.set_font('Arial', 'B', 12)
		pdf.cell(259.4,2,' ',0,1,'C')
		pdf.cell(259.4,2,' ',0,1,'C')
		pdf.cell(259.4,3,'"'+title+'"',0,1,'C')
		pdf.set_font('Arial', '', 12)
		pdf.cell(259.4,2,' ',0,1,'C')
		pdf.cell(259.4,2,' ',0,1,'C')
		pdf.cell(259.4,3,'conducted on '+date+' organized by',0,1,'C')
		pdf.cell(259.4,2,' ',0,1,'C')
		pdf.cell(259.4,2,' ',0,1,'C')
		pdf.cell(259.4,3,' '+org+' ',0,1,'C')
		pdf.cell(259.4,55,' ',0,1,'C')
		print(dep)
		if dep=='Computer':
			pdf.cell(259.4,5,' '+con1+'              Mrs. Saritha Ambedkar               Dr. Sunitha Patil              Dr. Suresh Ukarande',0,1,'C')
		elif dep=='IT':
			pdf.cell(259.4,5,' '+con1+'              Mrs. Vaishali Wadhe               Dr. Sunitha Patil              Dr. Suresh Ukarande',0,1,'C')
		elif dep=='EXTC':
			pdf.cell(259.4,5,' '+con1+'              Mrs. Aarthi Sahithya               Dr. Sunitha Patil              Dr. Suresh Ukarande',0,1,'C')
		elif dep=='All':
			pdf.cell(259.4,5,' '+con1+'                                                  Dr. Sunitha Patil              Dr. Suresh Ukarande',0,1,'C')
		else :
			pdf.cell(259.4,5,' '+con1+'              '+con2+'               Dr. Sunitha Patil              Dr. Suresh Ukarande',0,1,'C')
		pdf.set_font('Arial', '', 9)
		pdf.cell(259.4,6,' ( Convener of Event )                           ( Head of the Department )                   ( Vice Principal,KJSIEIT )                   ( Principal, KJSIEIT )',0,1,'C')
		file='Certificate_'+evid+'.pdf'
		response = HttpResponse(pdf.output(name='Certificate_'+evid+'.pdf',dest='S').encode('latin-1'), content_type='application/pdf')
		content = "attachment; filename=%s " %(file)
		response['Content-Disposition'] = content
		messages.success(request,"Downloaded Report to Certificate_"+evid+".pdf file in your desktop")
		return response
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')



def torep(request):
	if request.method == "POST":
		evid=request.POST.get("dell", "")
		rep=Addreport.objects.get(id=evid)
		rdate=str(rep.erref.date)
		rfrom=str(rep.erref.tfrom)
		rto=str(rep.erref.tto)
		rnum=str(rep.id)
		pdf = fpdf.FPDF(format='letter')
		pdf.add_page()
		pdf.set_text_color(167, 36, 41)
		pdf.set_fill_color(0, 250, 250)
		pdf.set_font('Arial', '', 13)
		name1='https://departmentalactivity.herokuapp.com/static/kj.png'
		name2='https://departmentalactivity.herokuapp.com/static/kj2.png'
		font1='https://departmentalactivity.herokuapp.com/static/algerianr.ttf'
		font2='https://departmentalactivity.herokuapp.com/static/calibb.ttf'
		pdf.cell(194,2,' ',0,1,'C')
		pdf.cell(194,5.5,'K.J. Somaiya Insititute of Engineering and',0,1,'C')
		pdf.image(name1, x = 20, y = 10, w = 24, h = 24, type = 'png', link = 'static')
		pdf.set_fill_color(0, 0, 250)
		pdf.cell(194,5.5,'Information Technology Sion, Mumbai',0,1,'C')
		pdf.cell(194,1,' ',0,1,'C')
		pdf.set_font('Times', 'B', 8)
		pdf.set_text_color(3, 90, 171)
		pdf.cell(194,3,'KJSIEIT IS ACCREDITED BY NAAC WITH "A" GRADE',0,1,'C')
		pdf.cell(194,3,'THREE PROGRAMS ACCREDITED BY NATIONAL BOARD OF ACCREDITATION',0,1,'C')
		pdf.cell(194,3,'BEST COLLEGE AWARDED BY UNIVERSITY OF MUMBAI, ISTE (MH) AND CSI (MUMBAI)',0,1,'C')
		pdf.image(name2, x = 175, y = 13, w = 24, h = 18, type = 'png', link = 'static')
		pdf.ln()
		pdf.line(10, 39, 205, 39)
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.ln()
		pdf.set_font('Arial', '', 14)
		pdf.set_text_color(0,0,0)
		pdf.set_font("Arial", size=19)
		pdf.set_text_color(0,0,0)
		pdf.cell(195,5,' ',0,1,'L')
		pdf.cell(195,11,'Report : '+rep.erref.title,0,1,'C')
		pdf.set_font('Arial', '', 11)
		pdf.set_text_color(100,100,100)
		pdf.set_font("Arial", size=12)
		pdf.set_text_color(30,30,30)
		pdf.cell(195,5,' ',0,1,'L')
		pdf.cell(195,6,'Date : '+rdate+' ',0,1,'R')
		pdf.cell(195,6,'Time : '+rfrom+' - '+rto,0,1,'R')
		pdf.set_font('Arial', '', 11)
		pdf.set_text_color(40,40,40)
		pdf.ln()
		pdf.ln()
		pdf.set_font('Arial', '', 11)
		pdf.set_fill_color(0,254,254)
		body1=' at K J Somaiya Insititute of Engineering and Information Technology (KJSIEIT) organized a webinar on '
		body2='.The sesson was conducted by '+rep.erref.organizer+'.'
		pdf.multi_cell(195,6,rep.erref.department+body1+'"'+rep.erref.title+'" on '+rdate+body2,0,1,'C')
		pdf.ln()
		pdf.cell(195,0,' ',0,1,'L')
		pdf.cell(195,3,' ',0,1,'L')
		pdf.ln()
		pdf.ln()
		pdf.set_font('Arial', 'B', 11)
		pdf.set_text_color(40,40,40)
		print(len(rep.body))
		if rep.objectives:
			pdf.cell(195,3,'Objectives : ',0,1,'L')
			pdf.ln()
			pdf.set_font('Arial', '', 11)
			pdf.set_fill_color(0,254,254)
			pdf.multi_cell(195,4,rep.objectives,0,1,'C')
			pdf.ln()
			pdf.cell(195,0,'',1,1,'L')
			pdf.cell(195,3,' ',0,1,'L')
			pdf.ln()
			pdf.ln()
		pdf.set_font('Arial', 'B', 11)
		if rep.keypoints:
			pdf.cell(195,3,'Keypoints : ',0,1,'L')
			pdf.ln()
			pdf.set_font('Arial', '', 11)
			pdf.set_fill_color(0,254,254)
			pdf.multi_cell(195,4,rep.keypoints,0,1,'C')
			pdf.ln()
			pdf.cell(195,0,'',1,1,'L')
			pdf.cell(195,3,' ',0,1,'L')
			pdf.ln()
			pdf.ln()
		pdf.set_font('Arial', 'B', 11)
		if rep.outcomes:
			pdf.cell(195,3,'Outcomes : ',0,1,'L')
			pdf.ln()
			pdf.set_font('Arial', '', 11)
			pdf.set_fill_color(0,254,254)
			pdf.multi_cell(195,4,rep.outcomes,0,1,'C')
			pdf.ln()
			pdf.cell(195,0,'',1,1,'L')
			pdf.cell(195,3,' ',0,1,'L')
			pdf.ln()
			pdf.ln()
		pdf.set_font('Arial', 'B', 11)
		if rep.feedback:
			pdf.cell(195,3,'Feedback : ',0,1,'L')
			pdf.ln()
			pdf.set_font('Arial', '', 11)
			pdf.set_fill_color(0,254,254)
			pdf.multi_cell(195,4,rep.feedback,0,1,'C')
			pdf.ln()
			pdf.cell(195,0,'',1,1,'L')
			pdf.cell(195,3,' ',0,1,'L')
			pdf.ln()
			pdf.ln()
		pdf.set_font('Arial', '', 11)
		pdf.set_text_color(0,0,0)
		pdf.ln()
		pdf.cell(195,4,"Faculty Co-ordinators :  "+rep.erref.teacher1+",  "+rep.erref.teacher2+",  "+rep.erref.teacher3,0,1,'L')
		pdf.ln()
		pdf.cell(195,4,"Conveners :  "+rep.erref.convener1+",  "+rep.erref.convener2,0,1,'L')
		pdf.ln()
		pdf.cell(195,4,"Contact :  "+rep.erref.contact1+",  "+rep.erref.contact2,0,1,'L')
		file='Report_'+rnum+'.pdf'
		response = HttpResponse(pdf.output(name='Report_'+rnum+'.pdf',dest='S').encode('latin-1'), content_type='application/pdf')
		content = "attachment; filename=%s " %(file)
		response['Content-Disposition'] = content
		messages.success(request,"Downloaded Report to Report_"+rnum+".pdf file in your desktop")
		return response
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')


def todocall(request):
	if request.method == "POST":
		evid=request.POST.get("dell", "")
		title=request.POST.get("delltitle", "")
		date=request.POST.get("delldate", "")
		desc=request.POST.get("delldesc", "")
		etype=request.POST.get("delltype", "")
		timef=request.POST.get("dellfrom", "")
		timet=request.POST.get("dellto", "")
		regf=request.POST.get("dellregf", "")
		regt=request.POST.get("dellregt", "")
		org=request.POST.get("dellorg", "")
		pre=request.POST.get("dellpresenter", "")
		pred=request.POST.get("dellpresenterd", "")
		dep=request.POST.get("delldep", "")
		per=request.POST.get("dellperks", "")
		cer=request.POST.get("dellcerti", "")
		faccod1=request.POST.get("dellfaccod1", "")
		faccod2=request.POST.get("dellfaccod2", "")
		faccod3=request.POST.get("dellfaccod3", "")
		conv1=request.POST.get("dellconv1", "")
		conv2=request.POST.get("dellconv2", "")
		cont1=request.POST.get("dellcont1", "")
		cont2=request.POST.get("dellcont2", "")
		edoc = Document()
		h=edoc.add_heading('Event No '+evid+' Details', 0)
		p = edoc.add_heading(org+' is organizing a '+etype+' on topic "'+title+'" in association with '+dep, 1)
		p = edoc.add_paragraph()
		p = edoc.add_paragraph(''+desc)
		p = edoc.add_paragraph('Date : '+date+', from '+timef+' to'+timet)
		p = edoc.add_paragraph()
		p = edoc.add_paragraph('Presenter Details : '+pre+', '+pred)
		p = edoc.add_paragraph('Perks : '+per+', '+cer)
		p = edoc.add_paragraph('Registrations from '+regf+' to '+regt)
		p = edoc.add_paragraph()
		p = edoc.add_paragraph('Faculty Co-ordinators : '+faccod1+', '+faccod2+', '+faccod3)
		p = edoc.add_paragraph('Conveners : '+conv1+', '+conv2)
		p = edoc.add_paragraph('Contacts : '+cont1+', '+cont2)
		if path.isfile('Event'+evid+'.docx')==True:
			messages.warning(request,"File Event"+evid+".doc already exists in your desktop")
		else:
			messages.success(request,"Downloaded data to Event"+evid+".doc file in your desktop")
			file='Event'+evid+'.docx'
			response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
			content = "attachment; filename=%s " %(file)
			response['Content-Disposition'] = content
			edoc.save(response)
		return response
	else:
		return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')


# ----------00000000000000000000000-----0000000-----00000000000000000000000--------------------------------------
# ----------00000000000000000000000-----0000000-----0000000000000000000000000------------------------------------
# ----------00000000000000000000000-----0000000-----00000000000000000000000000-----------------------------------
# ----------0000000---------------------0000000-----0000000-------------0000000----------------------------------
# ----------0000000---------------------0000000-----0000000-------------0000000----------------------------------
# ----------0000000---------------------0000000-----0000000-------------0000000----------------------------------
# ----------00000000000000000000000-----0000000-----0000000-------------0000000----------------------------------
# ----------00000000000000000000000-----0000000-----0000000-------------0000000----------------------------------
# ----------0000000---------------------0000000-----0000000-------------0000000----------------------------------
# ----------0000000---------------------0000000-----0000000-------------0000000----------------------------------
# ----------0000000---------------------0000000-----0000000-------------0000000----------------------------------
# ----------00000000000000000000000-----0000000-----00000000000000000000000000-----------------------------------
# ----------00000000000000000000000-----0000000-----0000000000000000000000000------------------------------------
# ----------00000000000000000000000-----0000000-----00000000000000000000000--------------------------------------




# def sms(request):
# 	account_sid = "AC6ae2ac6a1f4166ea1d956ab56b936857"
# 	auth_token  = "348ad31dd0df2d05af2dcbfae8d91783"
# 	client = Client(account_sid, auth_token)

# 	message = client.messages.create(
# 		to="+919819150055",
# 		from_="9136247408",
# 		body="Hello Sir")
# 	print(message.sid)
# 	return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')

# def sms(request):
#     user = 'vivekbilla345@gmail.com'
#     password = 'EngineeringHR'
#     voice = Voice()
#     voice.login(user, password)
#     number = '9819150055' # use these for command method
#     message = 'Message text: '
#     voice.send_sms(number, message)


def email(request):
	user=request.user
	subject = 'Registration Successful!'
	message = ' it  means a world to us '+user.first_name
	email_from = settings.EMAIL_HOST_USER
	recipient_list = ['vivekananda.b@somaiya.edu',user.email]
	send_mail( subject, message, email_from, recipient_list )
	messages.success(request,'Mail Sent Successfully')
	return redirect('kjsomaiyacollegeofengineeringandinformationtechnologyteachers-allevents')